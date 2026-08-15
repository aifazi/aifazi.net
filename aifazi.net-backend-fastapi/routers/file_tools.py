"""
routers/file_tools.py — All file processing tools (Foxit-grade, backend Python)
Libraries: PyMuPDF (fitz), Pillow, python-docx, openpyxl, reportlab
"""
import base64
import csv
import html
import io
import json
import re
import zipfile

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from dependencies import require_staff

router = APIRouter()

MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB hard cap

# ── Helpers ───────────────────────────────────────────────────────
def _pdf_stream(doc, name="output.pdf"):
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True); doc.close(); buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{name}"'})

def _bytes_stream(data: bytes, mime: str, name: str):
    return StreamingResponse(io.BytesIO(data), media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{name}"'})

def _zip_stream(files: list[tuple[str, bytes]], name="archive.zip"):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for fname, data in files: zf.writestr(fname, data)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{name}"'})

async def _read(f: UploadFile) -> bytes:
    data = await f.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, f"File too large (max {MAX_UPLOAD_BYTES // 1024 // 1024} MB)")
    return data

# ════════════════════════════════════════════════════════
# PDF BASIC TOOLS
# ════════════════════════════════════════════════════════

@router.post("/pdf/merge")
async def merge_pdf(files: list[UploadFile] = File(...), _: dict = Depends(require_staff)):
    import fitz
    if len(files) < 2: raise HTTPException(400, "Need at least 2 PDFs")
    out = fitz.open()
    for f in files:
        doc = fitz.open(stream=await _read(f), filetype="pdf")
        out.insert_pdf(doc); doc.close()
    return _pdf_stream(out, "merged.pdf")

@router.post("/pdf/split")
async def split_pdf(file: UploadFile = File(...), pages: str = Form("1"), _: dict = Depends(require_staff)):
    """pages: '1,3,5-7' — returns ZIP of individual PDFs"""
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    total = doc.page_count
    def parse_ranges(s):
        result = set()
        for part in s.split(','):
            part = part.strip()
            if '-' in part:
                a, b = part.split('-', 1)
                result.update(range(int(a)-1, int(b)))
            elif part.isdigit():
                result.add(int(part)-1)
        return sorted(p for p in result if 0 <= p < total)
    sel = parse_ranges(pages) if pages.strip() else list(range(total))
    if not sel: raise HTTPException(400, "No valid pages specified")
    zfiles = []
    for i, p in enumerate(sel):
        sub = fitz.open(); sub.insert_pdf(doc, from_page=p, to_page=p)
        buf = io.BytesIO(); sub.save(buf, garbage=4, deflate=True); sub.close()
        zfiles.append((f"page_{p+1:03d}.pdf", buf.getvalue()))
    doc.close()
    return _zip_stream(zfiles, "split_pages.zip")

@router.post("/pdf/compress")
async def compress_pdf(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    return _pdf_stream(doc, "compressed.pdf")

@router.post("/pdf/rotate")
async def rotate_pdf(file: UploadFile = File(...), angle: int = Form(90),
                     pages: str = Form("all"), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    sel = range(doc.page_count) if pages.strip().lower() == 'all' else \
          [int(p)-1 for p in pages.split(',') if p.strip().isdigit()]
    for i in sel:
        if 0 <= i < doc.page_count:
            doc[i].set_rotation((doc[i].rotation + angle) % 360)
    return _pdf_stream(doc, "rotated.pdf")

@router.post("/pdf/remove-pages")
async def remove_pages(file: UploadFile = File(...), pages: str = Form(...), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    idxs = sorted(set(int(p)-1 for p in pages.split(',') if p.strip().isdigit()
                       and 0 < int(p) <= doc.page_count), reverse=True)
    for i in idxs: doc.delete_page(i)
    return _pdf_stream(doc, "removed_pages.pdf")

@router.post("/pdf/watermark")
async def watermark_pdf(file: UploadFile = File(...), text: str = Form("CONFIDENTIAL"),
    opacity: float = Form(0.2), angle: int = Form(45), color: str = Form("#cc0000"),
    font_size: int = Form(48), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    h = color.lstrip('#').ljust(6, '0') if len(color.lstrip('#')) < 6 else color.lstrip('#')
    try: rgb = (int(h[:2],16)/255, int(h[2:4],16)/255, int(h[4:6],16)/255)
    except: rgb = (0.8, 0, 0)
    for page in doc:
        r = page.rect; cx, cy = r.width/2, r.height/2
        mat = fitz.Matrix(1,0,0,1,0,0).prerotate(angle)
        shape = page.new_shape()
        shape.draw_rect(r)
        shape.finish(fill=None, color=None, width=0)
        page.insert_text((cx-len(text)*font_size*0.25, cy), text,
            fontsize=font_size, color=rgb, rotate=angle,
            render_mode=3)
        # overlay with transparency via annotation
        tw = fitz.TextWriter(r, color=rgb, opacity=opacity)
        tw.append((cx-len(text)*font_size*0.3, cy), text, fontsize=font_size)
        tw.write_text(page, opacity=opacity, angle=angle)
    return _pdf_stream(doc, "watermarked.pdf")

@router.post("/pdf/page-numbers")
async def page_numbers(file: UploadFile = File(...), position: str = Form("bottom-center"),
    start: int = Form(1), prefix: str = Form(""), font_size: int = Form(11),
    color: str = Form("#000000"), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    h = color.lstrip('#').ljust(6, '0') if len(color.lstrip('#')) < 6 else color.lstrip('#')
    try: rgb = (int(h[:2],16)/255, int(h[2:4],16)/255, int(h[4:6],16)/255)
    except: rgb = (0,0,0)
    for i, page in enumerate(doc):
        r = page.rect; n = i + start
        label = f"{prefix}{n}"
        x = {'bottom-center': r.width/2 - len(label)*font_size*0.25,
             'bottom-right':  r.width - 60,
             'bottom-left':   30,
             'top-center':    r.width/2 - len(label)*font_size*0.25,
             'top-right':     r.width - 60,
             'top-left':      30}.get(position, r.width/2)
        y = {'top-center':15,'top-left':15,'top-right':15}.get(position, r.height - 15)
        page.insert_text((x, y), label, fontsize=font_size, color=rgb)
    return _pdf_stream(doc, "numbered.pdf")

@router.post("/pdf/images-to-pdf")
async def images_to_pdf(files: list[UploadFile] = File(...), _: dict = Depends(require_staff)):
    import fitz
    from PIL import Image as PILImage
    out = fitz.open()
    for f in files:
        raw = await _read(f)
        try:
            img = PILImage.open(io.BytesIO(raw))
            w, h = img.size; r = fitz.Rect(0, 0, w*0.75, h*0.75)
            page = out.new_page(width=r.width, height=r.height)
            page.insert_image(r, stream=raw)
        except Exception: pass
    return _pdf_stream(out, "images.pdf")

@router.post("/pdf/to-images")
async def pdf_to_images(file: UploadFile = File(...), scale: float = Form(1.5),
    fmt: str = Form("png"), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    scale = max(0.5, min(scale, 4.0)); fmt = fmt.lower()
    mime = "image/jpeg" if fmt == "jpg" else "image/png"
    zfiles = []
    for i in range(doc.page_count):
        pix = doc[i].get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        data = pix.tobytes(fmt if fmt in ("png","jpg") else "png")
        zfiles.append((f"page_{i+1:03d}.{fmt}", data))
    doc.close()
    return _zip_stream(zfiles, "pages.zip")

# ════════════════════════════════════════════════════════
# PDF ADVANCED TOOLS
# ════════════════════════════════════════════════════════

@router.post("/pdf/protect")
async def protect_pdf(file: UploadFile = File(...), password: str = Form(...),
    owner_password: str = Form(""), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    perm = fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY
    enc = fitz.PDF_ENCRYPT_AES_256
    buf = io.BytesIO()
    doc.save(buf, encryption=enc, user_pw=password,
             owner_pw=owner_password or password, permissions=perm, garbage=4, deflate=True)
    doc.close(); buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="protected.pdf"'})

@router.post("/pdf/unlock")
async def unlock_pdf(file: UploadFile = File(...), password: str = Form(""), _: dict = Depends(require_staff)):
    import fitz
    raw = await _read(file)
    doc = fitz.open(stream=raw, filetype="pdf")
    if doc.needs_pass:
        if not doc.authenticate(password):
            raise HTTPException(400, "Wrong password")
    return _pdf_stream(doc, "unlocked.pdf")

@router.post("/pdf/organize")
async def organize_pdf(file: UploadFile = File(...), order: str = Form(...), _: dict = Depends(require_staff)):
    """order: comma-separated 1-based page numbers e.g. '3,1,2'"""
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    idxs = [int(p)-1 for p in order.split(',') if p.strip().isdigit()
            and 0 < int(p) <= doc.page_count]
    if not idxs: raise HTTPException(400, "No valid page order")
    out = fitz.open()
    for i in idxs: out.insert_pdf(doc, from_page=i, to_page=i)
    doc.close()
    return _pdf_stream(out, "organized.pdf")

@router.post("/pdf/crop")
async def crop_pdf(file: UploadFile = File(...),
    top: float=Form(0), bottom: float=Form(0), left: float=Form(0), right: float=Form(0), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    for page in doc:
        r = page.rect
        page.set_cropbox(fitz.Rect(r.x0+left, r.y0+top, r.x1-right, r.y1-bottom))
    return _pdf_stream(doc, "cropped.pdf")

@router.post("/pdf/edit-meta")
async def edit_meta(file: UploadFile = File(...), title: str = Form(""),
    author: str = Form(""), subject: str = Form(""), keywords: str = Form(""),
    creator: str = Form(""), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    doc.set_metadata({"title":title,"author":author,"subject":subject,
                      "keywords":keywords,"creator":creator})
    return _pdf_stream(doc, "edited.pdf")

@router.post("/pdf/info")
async def pdf_info(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    info = {"page_count": doc.page_count, "metadata": doc.metadata or {},
            "pages": [{"w": round(doc[i].rect.width,1), "h": round(doc[i].rect.height,1),
                        "rotation": doc[i].rotation} for i in range(doc.page_count)]}
    doc.close()
    return info

@router.post("/pdf/grayscale")
async def grayscale_pdf(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    out = fitz.open()
    for page in doc:
        pix = page.get_pixmap(colorspace=fitz.csGRAY, alpha=False)
        imgpdf = fitz.open("pdf", pix.pdfocr_tobytes() if hasattr(pix,'pdfocr_tobytes')
                           else pix.tobytes("png"))
        out.insert_pdf(imgpdf); imgpdf.close()
    doc.close()
    return _pdf_stream(out, "grayscale.pdf")

@router.post("/pdf/header-footer")
async def header_footer(file: UploadFile = File(...), header: str = Form(""),
    footer: str = Form(""), font_size: int = Form(10), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    for i, page in enumerate(doc):
        r = page.rect
        if header:
            page.insert_text((r.width/2 - len(header)*font_size*0.25, 18),
                header.replace('{n}', str(i+1)).replace('{total}', str(doc.page_count)),
                fontsize=font_size, color=(0.2,0.2,0.2))
        if footer:
            page.insert_text((r.width/2 - len(footer)*font_size*0.25, r.height - 12),
                footer.replace('{n}', str(i+1)).replace('{total}', str(doc.page_count)),
                fontsize=font_size, color=(0.2,0.2,0.2))
    return _pdf_stream(doc, "header_footer.pdf")

@router.post("/pdf/flatten")
async def flatten_pdf(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    for page in doc:
        for annot in page.annots(): annot.update()
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True, no_new_id=True)
    doc.close(); buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="flattened.pdf"'})

@router.post("/pdf/sign")
async def sign_pdf(file: UploadFile = File(...), name: str = Form(""),
    sig_image: UploadFile | None = File(None),
    page: int = Form(0), x: float = Form(50), y: float = Form(700),
    width: float = Form(200), height: float = Form(60),
    _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    if page >= doc.page_count: raise HTTPException(400, "Invalid page")
    p = doc[page]; rect = fitz.Rect(x, y, x+width, y+height)
    if sig_image:
        raw = await _read(sig_image)
        p.insert_image(rect, stream=raw)
    elif name:
        shape = p.new_shape(); shape.draw_rect(rect)
        shape.finish(color=(0.2,0.2,0.2), width=0.5)
        shape.commit()
        p.insert_text((x+6, y+height/2+5), name, fontsize=min(height*0.45, 24),
                      color=(0.05, 0.1, 0.6), fontname="helv")
    return _pdf_stream(doc, "signed.pdf")

@router.post("/pdf/repair")
async def repair_pdf(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    try:
        doc = fitz.open(stream=await _read(file), filetype="pdf")
    except Exception as e:
        raise HTTPException(400, f"Cannot read PDF: {e}")
    return _pdf_stream(doc, "repaired.pdf")

@router.post("/pdf/to-word")
async def pdf_to_word(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    text_blocks = []
    for i in range(doc.page_count):
        text_blocks.append(f"--- Page {i+1} ---\n{doc[i].get_text()}\n")
    doc.close()
    full_text = "\n".join(text_blocks)
    # Build minimal DOCX-like RTF (works without python-docx)
    try:
        from docx import Document
        d = Document(); d.add_heading("Extracted from PDF", 0)
        for block in text_blocks:
            for line in block.split('\n'):
                d.add_paragraph(line)
        buf = io.BytesIO(); d.save(buf); buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="extracted.docx"'})
    except ImportError:
        return _bytes_stream(full_text.encode(), "text/plain", "extracted.txt")

@router.post("/pdf/to-excel")
async def pdf_to_excel(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    all_rows = [["Page", "Text"]]
    for i in range(doc.page_count):
        lines = doc[i].get_text().split('\n')
        for ln in lines:
            if ln.strip(): all_rows.append([i+1, ln.strip()])
    doc.close()
    try:
        import openpyxl; wb = openpyxl.Workbook(); ws = wb.active; ws.title = "PDF Text"
        for row in all_rows: ws.append(row)
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": 'attachment; filename="extracted.xlsx"'})
    except ImportError:
        rows_csv = "\n".join(",".join(str(c) for c in row) for row in all_rows)
        return _bytes_stream(rows_csv.encode(), "text/csv", "extracted.csv")

# ════════════════════════════════════════════════════════
# PDF CONVERSIONS
# ════════════════════════════════════════════════════════

@router.post("/pdf/to-jpg")
async def pdf_to_jpg(file: UploadFile = File(...), scale: float = Form(1.5), _: dict = Depends(require_staff)):
    import fitz
    doc = fitz.open(stream=await _read(file), filetype="pdf")
    scale = max(0.5, min(scale, 4.0))
    zfiles = [(f"page_{i+1:03d}.jpg",
               doc[i].get_pixmap(matrix=fitz.Matrix(scale,scale),alpha=False).tobytes("jpg"))
              for i in range(doc.page_count)]
    doc.close(); return _zip_stream(zfiles, "pages_jpg.zip")

@router.post("/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    raw = await _read(file)
    try:
        doc = fitz.open(stream=raw, filetype="docx")
        return _pdf_stream(doc, "converted.pdf")
    except Exception:
        raise HTTPException(500, "Word conversion failed — ensure file is valid DOCX")

@router.post("/convert/excel-to-pdf")
async def excel_to_pdf(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    import openpyxl
    raw = await _read(file)
    wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True)
    # Build HTML table then convert
    parts = []
    for ws in wb.worksheets:
        parts.append(f"<h2>{html.escape(ws.title)}</h2><table border='1' cellpadding='4'>")
        for row in ws.iter_rows(values_only=True):
            parts.append("<tr>" + "".join(f"<td>{html.escape(str(c or ''))}</td>" for c in row) + "</tr>")
        parts.append("</table><br/>")
    htm = "<html><body style='font-family:sans-serif;font-size:11px'>" + "".join(parts) + "</body></html>"
    doc = fitz.open(); page = doc.new_page()
    page.insert_htmlbox(page.rect, htm, css="* {font-size:9pt}")
    return _pdf_stream(doc, "excel.pdf")

@router.post("/convert/csv-to-pdf")
async def csv_to_pdf(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import fitz
    raw = (await _read(file)).decode('utf-8', errors='replace')
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    parts = ["<html><body style='font-family:monospace;font-size:9px'><table border='1' cellpadding='3' width='100%'>"]
    for i, row in enumerate(rows):
        tag = "th" if i == 0 else "td"
        parts.append("<tr>" + "".join(f"<{tag}>{html.escape(str(c))}</{tag}>" for c in row) + "</tr>")
    parts.append("</table></body></html>")
    doc = fitz.open(); page = doc.new_page(width=842, height=595)
    page.insert_htmlbox(page.rect, "".join(parts))
    return _pdf_stream(doc, "csv_table.pdf")

@router.post("/convert/jpg-to-pdf")
async def jpg_to_pdf(files: list[UploadFile] = File(...), _: dict = Depends(require_staff)):
    import fitz
    from PIL import Image as PILImage
    out = fitz.open()
    for f in files:
        raw = await _read(f)
        img = PILImage.open(io.BytesIO(raw)); w, h = img.size
        page = out.new_page(width=w*0.75, height=h*0.75)
        page.insert_image(page.rect, stream=raw)
    return _pdf_stream(out, "images.pdf")

@router.post("/convert/html-to-pdf")
async def html_to_pdf(file: UploadFile | None = File(None),
    html_content: str = Form(""), _: dict = Depends(require_staff)):
    import fitz
    raw = (await _read(file)).decode('utf-8','replace') if file else html_content
    if not raw.strip(): raise HTTPException(400, "No HTML content provided")
    doc = fitz.open(); page = doc.new_page(width=794, height=1123)
    page.insert_htmlbox(page.rect, raw)
    return _pdf_stream(doc, "page.pdf")

@router.post("/convert/text-to-pdf")
async def text_to_pdf(file: UploadFile | None=File(None), text: str=Form(""),
    font_size: int=Form(11), title: str=Form(""), _: dict = Depends(require_staff)):
    import fitz
    raw = (await _read(file)).decode('utf-8','replace') if file else text
    if not raw.strip(): raise HTTPException(400, "No text provided")
    doc = fitz.open()
    page = doc.new_page(); lines = raw.split('\n'); y = 50; line_h = font_size * 1.4
    if title: page.insert_text((40,30), title, fontsize=font_size+4, color=(0,0,0))
    for ln in lines:
        if y + line_h > page.rect.height - 40:
            page = doc.new_page(); y = 50
        page.insert_text((40, y), ln, fontsize=font_size, color=(0,0,0))
        y += line_h
    return _pdf_stream(doc, "text.pdf")

# ════════════════════════════════════════════════════════
# IMAGE TOOLS (Pillow)
# ════════════════════════════════════════════════════════

@router.post("/image/compress")
async def compress_image(file: UploadFile = File(...), quality: int = Form(75), _: dict = Depends(require_staff)):
    from PIL import Image as PILImage
    img = PILImage.open(io.BytesIO(await _read(file)))
    fmt = img.format or 'JPEG'; buf = io.BytesIO()
    if img.mode in ('RGBA','P') and fmt == 'JPEG': img = img.convert('RGB')
    img.save(buf, format=fmt, quality=max(10,min(quality,95)), optimize=True)
    buf.seek(0); ext = fmt.lower().replace('jpeg','jpg')
    mime = 'image/jpeg' if fmt=='JPEG' else f'image/{fmt.lower()}'
    return _bytes_stream(buf.read(), mime, f"compressed.{ext}")

@router.post("/image/resize")
async def resize_image(file: UploadFile = File(...),
    width: int = Form(0), height: int = Form(0), keep_ratio: bool = Form(True), _: dict = Depends(require_staff)):
    from PIL import Image as PILImage
    img = PILImage.open(io.BytesIO(await _read(file)))
    ow, oh = img.size
    if width and height and not keep_ratio:
        img = img.resize((width, height), PILImage.LANCZOS)
    elif width:
        h = int(oh * width / ow); img = img.resize((width, h), PILImage.LANCZOS)
    elif height:
        w = int(ow * height / oh); img = img.resize((w, height), PILImage.LANCZOS)
    fmt = img.format or 'PNG'; buf = io.BytesIO()
    img.save(buf, format=fmt); buf.seek(0)
    ext = fmt.lower().replace('jpeg','jpg')
    return _bytes_stream(buf.read(), f'image/{fmt.lower()}', f"resized.{ext}")

@router.post("/image/convert")
async def convert_image(file: UploadFile = File(...), to_format: str = Form("png"), _: dict = Depends(require_staff)):
    from PIL import Image as PILImage
    img = PILImage.open(io.BytesIO(await _read(file)))
    fmt = to_format.upper().replace('JPG','JPEG')
    if fmt == 'JPEG' and img.mode in ('RGBA','P'): img = img.convert('RGB')
    buf = io.BytesIO(); img.save(buf, format=fmt); buf.seek(0)
    ext = to_format.lower(); mime = f'image/{"jpeg" if ext=="jpg" else ext}'
    return _bytes_stream(buf.read(), mime, f"converted.{ext}")

@router.post("/image/flip")
async def flip_image(file: UploadFile = File(...), direction: str = Form("horizontal"), _: dict = Depends(require_staff)):
    from PIL import Image as PILImage
    img = PILImage.open(io.BytesIO(await _read(file)))
    img = img.transpose(PILImage.FLIP_LEFT_RIGHT if direction=='horizontal' else PILImage.FLIP_TOP_BOTTOM)
    fmt = img.format or 'PNG'; buf = io.BytesIO(); img.save(buf, format=fmt); buf.seek(0)
    return _bytes_stream(buf.read(), f'image/{fmt.lower()}', f"flipped.{fmt.lower()}")

@router.post("/image/watermark")
async def watermark_image(file: UploadFile = File(...), text: str = Form("SAMPLE"),
    opacity: int = Form(50), color: str = Form("#ffffff"), font_size: int = Form(40), _: dict = Depends(require_staff)):
    from PIL import Image as PILImage
    from PIL import ImageDraw, ImageFont
    img = PILImage.open(io.BytesIO(await _read(file))).convert("RGBA")
    overlay = PILImage.new("RGBA", img.size, (0,0,0,0))
    draw = ImageDraw.Draw(overlay)
    h = color.lstrip('#').ljust(6,'0')
    try: rgb = (int(h[:2],16), int(h[2:4],16), int(h[4:6],16), int(opacity*2.55))
    except: rgb = (255,255,255,128)
    w, ht = img.size
    try: font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", font_size)
    except: font = ImageFont.load_default()
    bbox = draw.textbbox((0,0), text, font=font)
    tw, th = bbox[2]-bbox[0], bbox[3]-bbox[1]
    draw.text(((w-tw)//2, (ht-th)//2), text, fill=rgb, font=font)
    out = PILImage.alpha_composite(img, overlay).convert("RGB")
    fmt = "JPEG"; buf = io.BytesIO(); out.save(buf, format=fmt, quality=92); buf.seek(0)
    return _bytes_stream(buf.read(), "image/jpeg", "watermarked.jpg")

@router.post("/image/ocr")
async def image_ocr(file: UploadFile = File(...), _: dict = Depends(require_staff)):

    import fitz
    raw = await _read(file)
    # Try PyMuPDF's built-in OCR (uses Tesseract if available)
    try:
        pix = fitz.Pixmap(raw)
        text = pix.pdfocr_tobytes()  # returns PDF bytes with text layer
        # Extract text from the OCR'd PDF
        doc = fitz.open("pdf", text)
        extracted = "\n".join(doc[i].get_text() for i in range(doc.page_count))
        doc.close()
        if extracted.strip():
            return {"text": extracted}
    except Exception: pass
    raise HTTPException(503, "No OCR engine available")

# ════════════════════════════════════════════════════════
# DOCUMENT TOOLS
# ════════════════════════════════════════════════════════

@router.post("/doc/docx-to-text")
async def docx_to_text(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    raw = await _read(file)
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs)
        return {"text": text, "paragraphs": len(doc.paragraphs)}
    except ImportError:
        import fitz
        doc = fitz.open(stream=raw, filetype="docx")
        text = "\n".join(doc[i].get_text() for i in range(doc.page_count))
        doc.close(); return {"text": text}

@router.post("/doc/xlsx-to-csv")
async def xlsx_to_csv(file: UploadFile = File(...), sheet: str = Form("0"), _: dict = Depends(require_staff)):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(await _read(file)), data_only=True)
    ws = wb.worksheets[int(sheet)] if sheet.isdigit() else wb[sheet]
    buf = io.StringIO()
    writer = csv.writer(buf)
    for row in ws.iter_rows(values_only=True): writer.writerow(row)
    return _bytes_stream(buf.getvalue().encode(), "text/csv", f"{ws.title}.csv")

@router.post("/doc/csv-to-xlsx")
async def csv_to_xlsx(file: UploadFile = File(...), _: dict = Depends(require_staff)):
    import openpyxl
    raw = (await _read(file)).decode('utf-8', errors='replace')
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Sheet1"
    for row in csv.reader(io.StringIO(raw)): ws.append(row)
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="converted.xlsx"'})

# ════════════════════════════════════════════════════════
# TEXT TOOLS
# ════════════════════════════════════════════════════════

@router.post("/text/stats")
async def text_stats(file: UploadFile | None=File(None), text: str=Form(""), _: dict = Depends(require_staff)):
    raw = (await _read(file)).decode('utf-8','replace') if file else text
    words = raw.split(); sentences = re.split(r'[.!?]+', raw)
    freq: dict = {}
    for w in words:
        w2 = re.sub(r'[^a-zA-Z]','',w.lower())
        if len(w2) > 3: freq[w2] = freq.get(w2,0)+1
    top = sorted(freq.items(), key=lambda x:-x[1])[:20]
    return {"chars":len(raw), "chars_no_spaces":len(raw.replace(' ','')),
            "words":len(words), "sentences":len([s for s in sentences if s.strip()]),
            "paragraphs":len([p for p in raw.split('\n\n') if p.strip()]),
            "reading_time_min": round(len(words)/200, 1),
            "top_words": [{"word":w,"count":c} for w,c in top]}

@router.post("/text/compare")
async def compare_text(file1: UploadFile=File(...), file2: UploadFile=File(...), _: dict = Depends(require_staff)):
    import difflib
    t1 = (await _read(file1)).decode('utf-8','replace').splitlines()
    t2 = (await _read(file2)).decode('utf-8','replace').splitlines()
    diff = list(difflib.unified_diff(t1, t2, fromfile=file1.filename or "File 1",
                                     tofile=file2.filename or "File 2", lineterm=''))
    html_diff = difflib.HtmlDiff().make_file(t1, t2,
        fromdesc=file1.filename or "File 1", todesc=file2.filename or "File 2")
    return {"diff": "\n".join(diff), "diff_html": html_diff,
            "added": sum(1 for l in diff if l.startswith('+')),
            "removed": sum(1 for l in diff if l.startswith('-'))}

@router.post("/text/base64-encode")
async def base64_encode(file: UploadFile | None=File(None), text: str=Form(""), _: dict = Depends(require_staff)):
    if file:
        raw = await _read(file)
        return {"encoded": base64.b64encode(raw).decode(), "original_size": len(raw)}
    return {"encoded": base64.b64encode(text.encode()).decode()}

@router.post("/text/base64-decode")
async def base64_decode(text: str=Form(...), _: dict = Depends(require_staff)):
    try:
        dec = base64.b64decode(text.strip())
        try: return {"decoded": dec.decode('utf-8'), "type": "text"}
        except: return {"decoded": base64.b64encode(dec).decode(), "type": "binary", "size": len(dec)}
    except Exception as e:
        raise HTTPException(400, f"Invalid base64: {e}")

@router.post("/text/json-format")
async def json_format(file: UploadFile | None=File(None), text: str=Form(""),
    indent: int=Form(2), minify: bool=Form(False), _: dict = Depends(require_staff)):
    raw = (await _read(file)).decode('utf-8','replace') if file else text
    try:
        obj = json.loads(raw)
        out = json.dumps(obj, separators=(',',':')) if minify else json.dumps(obj, indent=indent, ensure_ascii=False)
        return {"formatted": out, "valid": True, "size": len(out)}
    except json.JSONDecodeError as e:
        return {"valid": False, "error": str(e), "formatted": raw}
